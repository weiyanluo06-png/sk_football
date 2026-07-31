from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


GREEN = "1F7A4D"
GREEN_DARK = "0F3D2A"
GREEN_PALE = "EAF3ED"
CREAM = "F7F4ED"
PAPER = "FFFDF8"
INK = "17231D"
MUTED = "65746B"
GOLD = "D8A63A"
LINE = "D7EADF"
CODE_BG = "F2F5F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_heading(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def apply_font(run, name="Microsoft YaHei", size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text: str, base_size=10.5, base_color=INK) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|<https?://[^>]+>)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            apply_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            apply_font(run, size=base_size, bold=True, color=GREEN_DARK)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            apply_font(run, name="Consolas", size=base_size - 0.5, color=GREEN_DARK)
            run.font.highlight_color = None
        else:
            run = paragraph.add_run(token[1:-1])
            apply_font(run, size=base_size, color=GREEN)
            run.underline = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        apply_font(run, size=base_size, color=base_color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.22

    heading_specs = {
        "Title": (28, GREEN_DARK, "Microsoft YaHei", True, 14, 10),
        "Subtitle": (12, MUTED, "Microsoft YaHei", False, 4, 10),
        "Heading 1": (17, GREEN_DARK, "Microsoft YaHei", True, 14, 7),
        "Heading 2": (13.5, GREEN, "Microsoft YaHei", True, 11, 5),
        "Heading 3": (11.5, GREEN_DARK, "Microsoft YaHei", True, 8, 3),
    }
    for style_name, (size, color, font, bold, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(GREEN_DARK)
    code.paragraph_format.left_indent = Cm(0.35)
    code.paragraph_format.right_indent = Cm(0.25)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.05

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Microsoft YaHei"
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    callout.font.size = Pt(10)
    callout.font.color.rgb = RGBColor.from_string(GREEN_DARK)
    callout.paragraph_format.left_indent = Cm(0.45)
    callout.paragraph_format.right_indent = Cm(0.25)
    callout.paragraph_format.space_before = Pt(3)
    callout.paragraph_format.space_after = Pt(3)


def shade_paragraph(paragraph, fill: str, border_color: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        p_pr.append(borders)


def add_cover(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(60)
    run = paragraph.add_run("SK FOOTBALL  /  PROJECT HANDBOOK")
    apply_font(run, name="Oswald", size=10, bold=True, color=GREEN)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(12)
    add_inline(title, "生康足球队纪念站", base_size=28, base_color=GREEN_DARK)

    subtitle = doc.add_paragraph(style="Subtitle")
    add_inline(subtitle, "项目交接手册", base_size=18, base_color=GREEN)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(15)
    rule.paragraph_format.space_after = Pt(28)
    shade_paragraph(rule, GREEN)
    rule.add_run(" ")

    statement = doc.add_paragraph()
    statement.paragraph_format.space_after = Pt(28)
    run = statement.add_run("真实资料  ·  统一审美  ·  可持续维护")
    apply_font(run, size=12, bold=True, color=GREEN_DARK)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(12.5)
    for line in lines:
        content = line.lstrip("> ").rstrip("  ").strip()
        if "：" not in content:
            continue
        label, value = content.split("：", 1)
        row = table.add_row()
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(12.5)
        set_cell_shading(row.cells[0], GREEN_PALE)
        set_cell_shading(row.cells[1], PAPER)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        r0 = p0.add_run(label)
        apply_font(r0, size=9, bold=True, color=GREEN_DARK)
        add_inline(p1, value, base_size=9, base_color=INK)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(28)
    run = note.add_run("交接目标")
    apply_font(run, size=9, bold=True, color=GREEN)
    body = doc.add_paragraph()
    body.paragraph_format.space_after = Pt(0)
    add_inline(
        body,
        "让下一位维护者能读懂数据、更新内容、延续设计，并始终把球队的真实记忆放在第一位。",
        base_size=11,
        base_color=INK,
    )

    doc.add_page_break()


def add_contents(doc: Document, headings: list[tuple[int, str]]) -> None:
    title = doc.add_paragraph(style="Heading 1")
    title.add_run("目录").font.color.rgb = RGBColor.from_string(GREEN_DARK)
    intro = doc.add_paragraph()
    add_inline(intro, "以下章节按实际维护顺序组织，可作为日常更新检查表使用。", base_color=MUTED)
    for level, text in headings:
        if level != 2:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.space_after = Pt(3)
        number = text.split(".", 1)[0] if "." in text else ""
        if number.isdigit():
            num_run = p.add_run(number.zfill(2))
            apply_font(num_run, name="Oswald", size=9.5, bold=True, color=GREEN)
            sep = p.add_run("   ")
            apply_font(sep, size=9.5)
        add_inline(p, text, base_size=10.5, base_color=INK)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            fill = GREEN_DARK if row_index == 0 else (PAPER if row_index % 2 else GREEN_PALE)
            set_cell_shading(cell, fill)
            text = row[col_index] if col_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            add_inline(
                paragraph,
                text,
                base_size=8.3,
                base_color="FFFFFF" if row_index == 0 else INK,
            )
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, code_lines: list[str], language: str) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after = Pt(1)
    label.paragraph_format.keep_with_next = True
    run = label.add_run((language or "CODE").upper())
    apply_font(run, name="Oswald", size=7.5, bold=True, color=GREEN)
    for code_line in code_lines or [""]:
        p = doc.add_paragraph(style="Code Block")
        shade_paragraph(p, CODE_BG, GREEN)
        run = p.add_run(code_line or " ")
        apply_font(run, name="Consolas", size=8.5, color=GREEN_DARK)


def add_body(doc: Document, lines: list[str]) -> None:
    index = 0
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    skipped_front_matter = False
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_language = stripped[3:].strip()
                code_lines = []
            else:
                add_code_block(doc, code_lines, code_language)
                in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if not skipped_front_matter:
            if stripped.startswith("# "):
                index += 1
                continue
            if stripped.startswith(">"):
                index += 1
                continue
            if stripped == "---" or not stripped:
                index += 1
                if stripped == "---":
                    skipped_front_matter = True
                continue

        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(
                p,
                heading.group(2),
                base_size={1: 17, 2: 13.5, 3: 11.5}[min(level, 3)],
                base_color=GREEN_DARK if level != 2 else GREEN,
            )
            set_repeat_heading(p)
            index += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Callout")
            shade_paragraph(p, GREEN_PALE, GREEN)
            add_inline(p, stripped.lstrip("> ").rstrip("  "), base_size=10, base_color=GREEN_DARK)
            index += 1
            continue

        bullet = re.match(r"^(\s*)-\s+(.+)$", raw)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.65 + len(bullet.group(1)) * 0.12)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, bullet.group(2), base_size=10.2)
            index += 1
            continue

        numbered = re.match(r"^(\s*)(\d+)\.\s+(.+)$", raw)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.7 + len(numbered.group(1)) * 0.12)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, numbered.group(3), base_size=10.2)
            index += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped.rstrip("  "), base_size=10.5)
        index += 1


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.68)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.28)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run("生康足球队纪念站  /  PROJECT HANDOFF")
        apply_font(run, name="Oswald", size=7.5, bold=True, color=MUTED)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run("武汉纺织大学生物医学工程与健康学院足球队   ·   ")
        apply_font(run, size=7.5, color=MUTED)
        add_page_number(paragraph)


def set_document_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.title = "生康足球队纪念站项目交接手册"
    props.subject = "项目维护、数据同步、视觉规范和 GitHub Pages 发布交接"
    props.author = "生康足球队纪念站"
    props.keywords = "足球队, 项目交接, 数据同步, 前端设计, GitHub Pages"


def build(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    front_matter = [line for line in lines[:20] if line.strip().startswith(">")]
    headings: list[tuple[int, str]] = []
    for line in lines:
        match = re.match(r"^(#{2,4})\s+(.+)$", line.strip())
        if match:
            headings.append((len(match.group(1)), match.group(2)))

    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    set_document_metadata(doc)
    add_cover(doc, front_matter)
    add_contents(doc, headings)
    add_body(doc, lines)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the football site handoff DOCX from Markdown.")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
